from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "reports" / "project_reports" / "rank_optimized_nn5_strategy_report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_table(doc: Document, df: pd.DataFrame, widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                cells[i].text = f"{value:.4f}"
            else:
                cells[i].text = str(value)
    format_table(table, widths)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def metric(value: float, pct: bool = False) -> str:
    if pct:
        return f"{value * 100:.1f}%"
    return f"{value:.4f}"


def short_model_name(value: str) -> str:
    mapping = {
        "rank_signed_anti_crowded38": "NN5 38F",
        "rank_signed_anti_crowded38_ic_weighted_composite": "Composite 38F",
        "rank_signed_small_ic_blend46_ic_weighted_composite": "Composite 46F",
    }
    return mapping.get(str(value), str(value))


def short_strategy_name(value: str) -> str:
    mapping = {
        "top5_equal_long_short_vol0.1": "Top 5%, equal, vol 10%",
        "top5_value_long_short_vol0.15": "Top 5%, value, vol 15%",
        "top5_value_long_short_vol0.2": "Top 5%, value, vol 20%",
        "top10_signal_long_short_volnone": "Top 10%, signal, no vol",
        "top10_equal_long_short_volnone": "Top 10%, equal, no vol",
        "top5_signal_long_short_volnone": "Top 5%, signal, no vol",
        "top5_equal_long_short_volnone": "Top 5%, equal, no vol",
        "top10_equal_long_short_vol0.2": "Top 10%, equal, vol 20%",
        "top10_signal_long_short_vol0.2": "Top 10%, signal, vol 20%",
        "top15_signal_long_short_volnone": "Top 15%, signal, no vol",
        "top15_signal_long_short_vol0.2": "Top 15%, signal, vol 20%",
    }
    return mapping.get(str(value), str(value))


def main() -> None:
    nn5_summary = pd.read_csv(
        ROOT
        / "reports/feature_engineering/nn5_rank_optimized_feature_research_2002_2016_fixed_nn5_summary.csv"
    )
    composite_summary = pd.read_csv(
        ROOT
        / "reports/feature_engineering/nn5_rank_optimized_feature_research_2002_2016_fixed_composite_summary.csv"
    )
    best_strategy = pd.read_csv(
        ROOT
        / "reports/strategies/rank_optimized_existing_predictions_fast/best_strategy_by_model.csv"
    )
    feature_ic = pd.read_csv(
        ROOT
        / "reports/feature_engineering/nn5_rank_optimized_feature_research_2002_2016_fixed_signed_feature_oos_ic.csv"
    )
    grid = pd.read_csv(
        ROOT
        / "reports/strategies/rank_optimized_existing_predictions_fast/all_strategy_grid_results.csv"
    )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Rank-Optimized NN5 Feature Research and Strategy Construction")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "Project report based on existing model-panel, saved NN5 forecasts, and strategy outputs"
    ).italic = True

    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph()
    p.add_run("Main result. ").bold = True
    p.add_run(
        "The rank-focused research improved the stock-ranking signal and converted it into a stronger portfolio result. "
        "The final NN5 rank-optimized forecast over 2002-2016 achieved monthly Spearman IC of "
        f"{nn5_summary.loc[0, 'monthly_spearman_ic']:.4f}. "
        "After strategy optimization, the selected NN5 long-short strategy reached full-period Sharpe "
        f"{best_strategy.iloc[0]['full_annualized_sharpe']:.4f} and post-tuning test Sharpe "
        f"{best_strategy.iloc[0]['test_annualized_sharpe']:.4f}."
    )
    add_bullet(
        doc,
        "Higher IC means better cross-sectional ranking, but the portfolio still depends on how the top and bottom groups are formed and weighted.",
    )
    add_bullet(
        doc,
        "The strongest strategy rule selected on 2002-2009 was top/bottom 5%, equal-weighted long-short, with 10% volatility targeting.",
    )
    add_bullet(
        doc,
        "The signed IC-weighted composite had higher raw IC than NN5, but the optimized NN5 portfolio produced the strongest selected post-tune Sharpe.",
    )

    doc.add_heading("Data and Inputs", level=1)
    add_bullet(doc, "Model panel: data/processed/model_penal_gkx_clean_rankfix.parquet.")
    add_bullet(
        doc,
        "Feature-selection ranks: 1987-1996 selection-window IC file from the feature research.",
    )
    add_bullet(doc, "Rank-optimized NN5 OOS prediction file: 2002-2016, 38 signed features.")
    add_bullet(
        doc,
        "Strategy optimization used only saved forecasts; no model retraining was performed in the strategy step.",
    )

    doc.add_heading("Research Workflow", level=1)
    steps = [
        (
            "1. Diagnose feature IC",
            "Measure each feature's monthly cross-sectional Spearman IC against next-month excess return.",
        ),
        (
            "2. Align signs",
            "Features with negative selection IC were multiplied by -1 so higher signed feature values generally point to higher expected return ranks.",
        ),
        (
            "3. Build compact feature sets",
            "Construct economically mixed sets under 50 variables, including risk/liquidity, value, profitability, momentum, accounting, intangible, growth, macro, and industry signals.",
        ),
        (
            "4. Change NN5 selection objective",
            "Keep MSE training loss, but select checkpoints and validation-grid candidates by validation Spearman IC instead of validation OOS R2.",
        ),
        (
            "5. Test OOS",
            "Run the expensive NN5 over the last 15 years, 2002-2016, and compare to signed IC-weighted composite forecasts.",
        ),
        (
            "6. Optimize strategy rules",
            "Use 2002-2009 to select top fraction, weighting, and volatility targeting; evaluate the selected rule on 2010-2016.",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Step"
    table.rows[0].cells[1].text = "What it did"
    for step, desc in steps:
        cells = table.add_row().cells
        cells[0].text = step
        cells[1].text = desc
    format_table(table, [1.75, 4.75])

    doc.add_heading("Feature IC Evidence", level=1)
    doc.add_paragraph(
        "The strongest signed OOS feature ICs came from liquidity/risk, value, and profitability-quality variables. "
        "Because signs were aligned, the ICs below are positive after transformation."
    )
    top_features = (
        feature_ic[feature_ic["model"] == "rank_signed_anti_crowded38"]
        .sort_values("importance_rank")
        .head(12)[["feature", "category", "mean_monthly_spearman_ic", "months"]]
        .rename(
            columns={
                "feature": "Feature",
                "category": "Category",
                "mean_monthly_spearman_ic": "OOS IC",
                "months": "Months",
            }
        )
    )
    add_table(doc, top_features, [1.25, 2.2, 1.1, 0.9])

    doc.add_heading("Forecast Results", level=1)
    forecast_table = pd.concat(
        [
            nn5_summary.assign(source="Rank-optimized NN5"),
            composite_summary.assign(source="Signed composite"),
        ],
        ignore_index=True,
    )[["source", "model", "n_features", "monthly_spearman_ic", "decile_10_minus_1_sharpe"]].rename(
        columns={
            "source": "Source",
            "model": "Model",
            "n_features": "Features",
            "monthly_spearman_ic": "Monthly IC",
            "decile_10_minus_1_sharpe": "Simple Decile Sharpe",
        }
    )
    forecast_table["Model"] = forecast_table["Model"].map(short_model_name)
    add_table(doc, forecast_table, [1.45, 2.5, 0.75, 0.85, 1.2])

    doc.add_heading("Strategy Optimization", level=1)
    doc.add_paragraph(
        "The strategy optimizer tested long-short rules built from saved predictions. "
        "It selected the rule on 2002-2009 using a Sharpe score penalized for drawdown and volatility, then evaluated on 2010-2016."
    )
    strategy_table = best_strategy[
        [
            "model",
            "strategy",
            "tune_annualized_sharpe",
            "test_annualized_sharpe",
            "full_annualized_sharpe",
            "full_max_drawdown",
        ]
    ].rename(
        columns={
            "model": "Forecast",
            "strategy": "Selected Strategy",
            "tune_annualized_sharpe": "Tune Sharpe",
            "test_annualized_sharpe": "Test Sharpe",
            "full_annualized_sharpe": "Full Sharpe",
            "full_max_drawdown": "Full Max DD",
        }
    )
    strategy_table["Forecast"] = strategy_table["Forecast"].map(short_model_name)
    strategy_table["Selected Strategy"] = strategy_table["Selected Strategy"].map(
        short_strategy_name
    )
    add_table(doc, strategy_table, [1.65, 1.75, 0.8, 0.8, 0.8, 0.8])

    doc.add_heading("Best Test-Period Rules", level=2)
    top_grid = (
        grid.sort_values("test_annualized_sharpe", ascending=False)
        .head(8)[
            [
                "model",
                "strategy",
                "tune_annualized_sharpe",
                "test_annualized_sharpe",
                "full_annualized_sharpe",
                "full_max_drawdown",
            ]
        ]
        .rename(
            columns={
                "model": "Forecast",
                "strategy": "Strategy",
                "tune_annualized_sharpe": "Tune Sharpe",
                "test_annualized_sharpe": "Test Sharpe",
                "full_annualized_sharpe": "Full Sharpe",
                "full_max_drawdown": "Full Max DD",
            }
        )
    )
    top_grid["Forecast"] = top_grid["Forecast"].map(short_model_name)
    top_grid["Strategy"] = top_grid["Strategy"].map(short_strategy_name)
    add_table(doc, top_grid, [1.55, 1.85, 0.75, 0.75, 0.75, 0.75])

    doc.add_heading("Interpretation", level=1)
    add_bullet(
        doc,
        "The raw signed composite has very high IC because it is directly built from signed individual signals; however, its OOS R2 is not meaningful because the composite score is a rank score, not a return-level forecast.",
    )
    add_bullet(
        doc,
        "The NN5 forecast has lower IC than the composite but a valid return-level scale and positive pooled OOS R2, which makes it easier to combine with volatility targeting.",
    )
    add_bullet(
        doc,
        "The best practical improvement came from reducing the portfolio to the strongest forecast tails and controlling realized volatility.",
    )

    doc.add_heading("Caveats and Next Steps", level=1)
    add_bullet(
        doc,
        "The strategy optimizer is still small and should be treated as research, not a production trading rule.",
    )
    add_bullet(
        doc,
        "Transaction costs, borrow costs, short-sale constraints, and turnover constraints are not yet included.",
    )
    add_bullet(
        doc,
        "A stronger next version should add turnover-aware weighting and a nested rolling strategy-selection process.",
    )

    doc.add_heading("Key Output Files", level=1)
    files = [
        "reports/feature_engineering/nn5_rank_optimized_feature_research_2002_2016_fixed_nn5_summary.csv",
        "reports/feature_engineering/nn5_rank_optimized_feature_research_2002_2016_fixed_nn5_predictions.parquet",
        "reports/feature_engineering/nn5_rank_optimized_feature_research_2002_2016_fixed_signed_feature_oos_ic.csv",
        "reports/strategies/rank_optimized_existing_predictions_fast/best_strategy_by_model.csv",
        "strategy/optimize_rank_strategy_fast.py",
    ]
    for item in files:
        add_bullet(doc, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
